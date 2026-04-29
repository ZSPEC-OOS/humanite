"""
Tests for batch processing, preset management, and document export.
All upstream services mocked. PostgreSQL mocked via app.dependency_overrides.
"""
import contextlib
import hashlib
import io
import os
from datetime import datetime, timezone
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

os.environ.setdefault("DATABASE_URL", "postgresql+asyncpg://test:test@localhost/test")
os.environ.setdefault("ENVIRONMENT",  "test")

SAMPLE_TEXT_A = (
    "Furthermore, Apple Inc. reported revenue of $89.5 billion in Q3 2024. "
    "The Federal Reserve does not expect rate cuts before Q2 2025."
)
SAMPLE_TEXT_B = (
    "Moreover, global GDP grew by 3.2% in fiscal year 2024, driven by "
    "strong consumer spending and robust manufacturing output worldwide."
)

AUTH_HEADERS = {"X-User-ID": "user-test-123", "X-User-Tier": "pro"}


# ── DB mock helpers ───────────────────────────────────────────────────────────

def _mock_db_no_existing_job():
    db = AsyncMock()
    db.add     = MagicMock()
    db.commit  = AsyncMock()
    db.delete  = AsyncMock()
    db.rollback = AsyncMock()
    mock_result = MagicMock()
    mock_result.scalar_one_or_none.return_value = None
    mock_result.scalars.return_value.all.return_value = []
    db.execute = AsyncMock(return_value=mock_result)
    db.get     = AsyncMock(return_value=None)
    return db


def _mock_db_with_existing_job():
    existing = MagicMock()
    existing.id     = "existing-job-uuid-000"
    existing.status = "completed"

    db = AsyncMock()
    db.add     = MagicMock()
    db.commit  = AsyncMock()
    db.rollback = AsyncMock()

    mock_dedup = MagicMock()
    mock_dedup.scalar_one_or_none.return_value = existing

    call_count = [0]

    async def execute_side_effect(query, *args, **kwargs):
        call_count[0] += 1
        if call_count[0] == 1:
            return mock_dedup
        empty = MagicMock()
        empty.scalar_one_or_none.return_value = None
        empty.scalars.return_value.all.return_value = []
        return empty

    db.execute = AsyncMock(side_effect=execute_side_effect)
    db.get     = AsyncMock(return_value=None)
    return db


@contextlib.contextmanager
def _override_db(mock_db):
    from src.main import app
    from src.database import get_db

    async def _mock_get_db():
        yield mock_db

    app.dependency_overrides[get_db] = _mock_get_db
    try:
        yield
    finally:
        app.dependency_overrides.pop(get_db, None)


# ── Batch endpoint ─────────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_batch_submit_returns_202_with_batch_id():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    with _override_db(_mock_db_no_existing_job()), \
         patch("src.routers.batch.queue_batch") as mock_q:
        mock_q.delay = MagicMock()

        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test",
            headers=AUTH_HEADERS,
        ) as client:
            resp = await client.post("/v1/batch", json={
                "items": [
                    {"item_id": "a", "text": SAMPLE_TEXT_A, "operation": "humanize"},
                    {"item_id": "b", "text": SAMPLE_TEXT_B, "operation": "humanize"},
                ]
            })

    assert resp.status_code == 202
    data = resp.json()
    assert "batch_job_id" in data
    assert data["status"] == "pending"
    assert data["total_items"] == 2
    assert data["accepted_items"] == 2
    assert data["rejected_items"] == 0
    assert data["poll_url"].startswith("/v1/batch/jobs/")
    mock_q.delay.assert_called_once()


@pytest.mark.asyncio
async def test_batch_rejects_items_under_min_length():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    with _override_db(_mock_db_no_existing_job()), \
         patch("src.routers.batch.queue_batch") as mock_q:
        mock_q.delay = MagicMock()

        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test",
            headers=AUTH_HEADERS,
        ) as client:
            resp = await client.post("/v1/batch", json={
                "items": [
                    {"item_id": "good",  "text": SAMPLE_TEXT_A, "operation": "humanize"},
                    {"item_id": "short", "text": "Too short.",   "operation": "humanize"},
                ]
            })

    assert resp.status_code == 202
    data = resp.json()
    assert data["accepted_items"] == 1
    assert data["rejected_items"] == 1

    statuses = {s["item_id"]: s for s in data["item_statuses"]}
    assert statuses["good"]["status"]  == "pending"
    assert statuses["short"]["status"] == "rejected"
    assert statuses["short"]["error_code"] == "VALIDATION_MIN_LENGTH"


@pytest.mark.asyncio
async def test_batch_exceeds_max_items_returns_400():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    items = [
        {"item_id": str(i), "text": SAMPLE_TEXT_A, "operation": "humanize"}
        for i in range(11)
    ]

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test",
        headers=AUTH_HEADERS,
    ) as client:
        resp = await client.post("/v1/batch", json={"items": items})

    assert resp.status_code == 400
    assert resp.json()["detail"]["code"] == "BATCH_ITEM_LIMIT_EXCEEDED"


@pytest.mark.asyncio
async def test_batch_duplicate_item_ids_rejected():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    with _override_db(_mock_db_no_existing_job()), \
         patch("src.routers.batch.queue_batch") as mock_q:
        mock_q.delay = MagicMock()

        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test",
            headers=AUTH_HEADERS,
        ) as client:
            resp = await client.post("/v1/batch", json={
                "items": [
                    {"item_id": "dup", "text": SAMPLE_TEXT_A, "operation": "humanize"},
                    {"item_id": "dup", "text": SAMPLE_TEXT_B, "operation": "humanize"},
                ]
            })

    data = resp.json()
    assert data["rejected_items"] == 1
    dup_statuses = [s for s in data["item_statuses"] if s["item_id"] == "dup"]
    assert any(s["error_code"] == "DUPLICATE_ITEM_ID" for s in dup_statuses)


@pytest.mark.asyncio
async def test_batch_deduplication_skips_existing_result():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    with _override_db(_mock_db_with_existing_job()), \
         patch("src.routers.batch.queue_batch") as mock_q:
        mock_q.delay = MagicMock()

        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test",
            headers=AUTH_HEADERS,
        ) as client:
            resp = await client.post("/v1/batch", json={
                "items": [
                    {"item_id": "existing", "text": SAMPLE_TEXT_A, "operation": "humanize"},
                ]
            })

    data = resp.json()
    statuses = {s["item_id"]: s for s in data["item_statuses"]}
    assert statuses["existing"]["status"] == "skipped"
    assert statuses["existing"]["job_id"] == "existing-job-uuid-000"


# ── Preset endpoints ───────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_create_preset_returns_201():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    with _override_db(_mock_db_no_existing_job()):
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test",
            headers=AUTH_HEADERS,
        ) as client:
            resp = await client.post("/v1/user/presets", json={
                "name":    "My Academic Preset",
                "intensity": 7,
                "tone":    "academic",
                "domain":  "academic",
                "preserve_citations": True,
            })

    assert resp.status_code == 201
    data = resp.json()
    assert data["name"] == "My Academic Preset"
    assert data["intensity"] == 7
    assert "id" in data
    assert "created_at" in data


@pytest.mark.asyncio
async def test_list_presets_returns_user_presets():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    mock_preset = MagicMock()
    mock_preset.id                 = "preset-uuid-001"
    mock_preset.name               = "Quick Clean"
    mock_preset.intensity          = 3
    mock_preset.tone               = "balanced"
    mock_preset.domain             = "general"
    mock_preset.preserve_citations = True
    mock_preset.created_at         = datetime.now(timezone.utc)

    mock_db = _mock_db_no_existing_job()
    mock_result = MagicMock()
    mock_result.scalars.return_value.all.return_value = [mock_preset]
    mock_db.execute = AsyncMock(return_value=mock_result)

    with _override_db(mock_db):
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test",
            headers=AUTH_HEADERS,
        ) as client:
            resp = await client.get("/v1/user/presets")

    assert resp.status_code == 200
    data = resp.json()
    assert isinstance(data, list)
    assert data[0]["name"] == "Quick Clean"
    assert data[0]["intensity"] == 3


@pytest.mark.asyncio
async def test_duplicate_preset_name_returns_409():
    from httpx import AsyncClient, ASGITransport
    from src.main import app
    from sqlalchemy.exc import IntegrityError

    mock_db = _mock_db_no_existing_job()
    mock_db.commit   = AsyncMock(side_effect=IntegrityError("", {}, None))
    mock_db.rollback = AsyncMock()

    with _override_db(mock_db):
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test",
            headers=AUTH_HEADERS,
        ) as client:
            resp = await client.post("/v1/user/presets", json={
                "name": "Existing Name", "intensity": 5,
                "tone": "balanced", "domain": "general",
            })

    assert resp.status_code == 409
    assert resp.json()["detail"]["code"] == "PRESET_NAME_TAKEN"


@pytest.mark.asyncio
async def test_delete_preset_returns_204():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    mock_preset = MagicMock()
    mock_preset.id      = "preset-uuid-del"
    mock_preset.user_id = "user-test-123"

    mock_db = _mock_db_no_existing_job()
    mock_db.get    = AsyncMock(return_value=mock_preset)
    mock_db.delete = AsyncMock()

    with _override_db(mock_db):
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test",
            headers=AUTH_HEADERS,
        ) as client:
            resp = await client.delete("/v1/user/presets/preset-uuid-del")

    assert resp.status_code == 204


# ── Export endpoints ───────────────────────────────────────────────────────────

EXPORT_PAYLOAD = {
    "text": (
        "Apple Inc. reported revenue of $89.5 billion in Q3 2024. "
        "The Federal Reserve does not expect rate cuts before Q2 2025."
    ),
    "format":    "text",
    "job_id":    "job-export-test-001",
    "title":     "Test Export",
    "watermark": {
        "type":             "ai_processed",
        "fingerprint":      "a" * 64,
        "job_id":           "job-export-test-001",
        "verification_url": "https://api.humanite.ai/v1/verify/" + "a" * 64,
        "issued_at":        "2024-01-01T00:00:00+00:00",
    },
}


@pytest.mark.asyncio
async def test_export_text_returns_plaintext_file():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test",
        headers=AUTH_HEADERS,
    ) as client:
        resp = await client.post("/v1/export", json=EXPORT_PAYLOAD)

    assert resp.status_code == 200
    assert "text/plain" in resp.headers["content-type"]
    assert "attachment" in resp.headers["content-disposition"]
    assert ".txt" in resp.headers["content-disposition"]

    body = resp.text
    assert "Apple Inc." in body
    assert "HUMANITE AI PROCESSING CERTIFICATE" in body
    assert "job-export-test-001" in body
    assert "a" * 64 in body


@pytest.mark.asyncio
async def test_export_markdown_contains_certificate_table():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    payload = {**EXPORT_PAYLOAD, "format": "markdown"}
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test",
        headers=AUTH_HEADERS,
    ) as client:
        resp = await client.post("/v1/export", json=payload)

    assert resp.status_code == 200
    assert "text/markdown" in resp.headers["content-type"]
    body = resp.text
    assert "# Test Export" in body
    assert "## Humanite AI Processing Certificate" in body
    assert "| Job ID |" in body
    assert "job-export-test-001" in body


@pytest.mark.asyncio
async def test_export_docx_returns_valid_docx_bytes():
    from httpx import AsyncClient, ASGITransport
    from src.main import app
    from docx import Document

    payload = {**EXPORT_PAYLOAD, "format": "docx"}
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test",
        headers=AUTH_HEADERS,
    ) as client:
        resp = await client.post("/v1/export", json=payload)

    assert resp.status_code == 200
    assert "wordprocessingml" in resp.headers["content-type"]
    assert ".docx" in resp.headers["content-disposition"]

    doc = Document(io.BytesIO(resp.content))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "Apple Inc." in all_text
    assert "Humanite AI Processing Certificate" in all_text


@pytest.mark.asyncio
async def test_export_docx_watermark_certificate_always_present():
    """Watermark certificate page must be present regardless of input content."""
    from httpx import AsyncClient, ASGITransport
    from src.main import app
    from docx import Document

    short_payload = {
        **EXPORT_PAYLOAD,
        "format": "docx",
        "text": "A minimal piece of text.",
    }
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test",
        headers=AUTH_HEADERS,
    ) as client:
        resp = await client.post("/v1/export", json=short_payload)

    doc = Document(io.BytesIO(resp.content))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "Humanite AI Processing Certificate" in all_text
    assert "job-export-test-001" in all_text


@pytest.mark.asyncio
async def test_export_invalid_format_returns_422():
    from httpx import AsyncClient, ASGITransport
    from src.main import app

    payload = {**EXPORT_PAYLOAD, "format": "pdf"}
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test",
        headers=AUTH_HEADERS,
    ) as client:
        resp = await client.post("/v1/export", json=payload)

    assert resp.status_code == 422
