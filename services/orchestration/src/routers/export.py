"""
Document export router.
Every exported document includes a mandatory watermark certificate that
cannot be disabled regardless of format or request parameters.
"""
import io
import logging
from datetime import datetime, timezone

from fastapi import APIRouter, Depends
from fastapi.responses import Response

from ..middleware.auth import get_user_claims
from ..schemas.export import ExportRequest

logger = logging.getLogger(__name__)
router = APIRouter()

_MIME_TYPES = {
    "text":     "text/plain; charset=utf-8",
    "markdown": "text/markdown; charset=utf-8",
    "docx":     "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

_EXTENSIONS = {
    "text":     "txt",
    "markdown": "md",
    "docx":     "docx",
}


@router.post("/v1/export")
async def export_document(
    body: ExportRequest,
    user: dict = Depends(get_user_claims),
) -> Response:

    job_id      = body.job_id
    fingerprint = body.watermark.get("fingerprint", "unavailable")
    verify_url  = body.watermark.get("verification_url", "https://api.humanite.ai/v1/verify")
    issued_at   = body.watermark.get("issued_at", datetime.now(timezone.utc).isoformat())
    fmt         = body.format
    ext         = _EXTENSIONS[fmt]
    filename    = f"humanite-{job_id[:8]}.{ext}"

    logger.info(
        "Export requested user_id=%s format=%s job_id=%s",
        user["user_id"], fmt, job_id,
    )

    if fmt == "text":
        content = _build_text(body.text, job_id, fingerprint, verify_url, issued_at)
        return Response(
            content=content.encode("utf-8"),
            media_type=_MIME_TYPES["text"],
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if fmt == "markdown":
        content = _build_markdown(body.text, body.title, job_id, fingerprint, verify_url, issued_at)
        return Response(
            content=content.encode("utf-8"),
            media_type=_MIME_TYPES["markdown"],
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    docx_bytes = _build_docx(body.text, body.title, job_id, fingerprint, verify_url, issued_at)
    return Response(
        content=docx_bytes,
        media_type=_MIME_TYPES["docx"],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _build_text(
    text: str,
    job_id: str,
    fingerprint: str,
    verify_url: str,
    issued_at: str,
) -> str:
    certificate = (
        "\n\n"
        + "─" * 60 + "\n"
        + "HUMANITE AI PROCESSING CERTIFICATE\n"
        + "─" * 60 + "\n"
        + f"Job ID:       {job_id}\n"
        + f"Fingerprint:  {fingerprint}\n"
        + f"Verify at:    {verify_url}\n"
        + f"Issued:       {issued_at}\n"
        + "This document was AI-generated text that has been processed\n"
        + "by Humanite AI. Original facts and entities are preserved.\n"
        + "─" * 60
    )
    return text.strip() + certificate


def _build_markdown(
    text: str,
    title: str,
    job_id: str,
    fingerprint: str,
    verify_url: str,
    issued_at: str,
) -> str:
    certificate = (
        "\n\n---\n\n"
        "## Humanite AI Processing Certificate\n\n"
        f"| Field | Value |\n"
        f"|---|---|\n"
        f"| Job ID | `{job_id}` |\n"
        f"| Fingerprint | `{fingerprint[:32]}…` |\n"
        f"| Verify | [{verify_url}]({verify_url}) |\n"
        f"| Issued | {issued_at} |\n\n"
        "_This document was AI-generated text processed by Humanite AI. "
        "Original facts and entities are preserved verbatim._\n"
    )
    return f"# {title}\n\n" + text.strip() + certificate


def _build_docx(
    text: str,
    title: str,
    job_id: str,
    fingerprint: str,
    verify_url: str,
    issued_at: str,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraphs = text.strip().split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph(para_text)
        p.style.font.size = Pt(11)

    doc.add_page_break()

    cert_heading = doc.add_heading("Humanite AI Processing Certificate", level=2)
    cert_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add job_id as a paragraph so it's discoverable via doc.paragraphs
    ref_para = doc.add_paragraph(f"Job ID: {job_id} | Verify: {verify_url}")
    ref_para.style.font.size = Pt(9)

    fields = [
        ("Job ID",      job_id),
        ("Fingerprint", fingerprint),
        ("Verify",      verify_url),
        ("Issued",      issued_at),
        ("Notice",
         "This document contains AI-generated text that has been processed "
         "by Humanite AI. All original facts, named entities, numbers, "
         "dates, and citations are preserved verbatim in the output."),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        label_run = row.cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)

        value_run = row.cells[1].paragraphs[0].add_run(value)
        value_run.font.size = Pt(9)
        if label == "Fingerprint":
            value_run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
